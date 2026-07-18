# -*- coding: utf-8 -*-
"""Renders content.DOCUMENT into a .docx engineering reference guide."""
import os
import re
import sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(__file__))
from content import DOCUMENT, ASSETS

HERE = os.path.dirname(__file__)
ASSETS_DIR = os.path.join(os.path.dirname(HERE), ASSETS)
OUT_PATH = os.path.join(os.path.dirname(HERE), "Engineer_Guide_Retail_SQL_Agent.docx")

INK = RGBColor(0x1A, 0x1F, 0x2B)
BLUE = RGBColor(0x1E, 0x40, 0xAF)
EMERALD = RGBColor(0x05, 0x76, 0x69)
RED = RGBColor(0x99, 0x1B, 0x1B)
AMBER = RGBColor(0x92, 0x40, 0x0E)
PURPLE = RGBColor(0x5B, 0x21, 0xB6)
MUTED = RGBColor(0x5B, 0x64, 0x72)

CALLOUT_FILL = {"note": "EFF6FF", "gotcha": "FFFBEB", "design": "F5F3FF", "safety": "FEF2F2"}
CALLOUT_BAR = {"note": "2563EB", "gotcha": "D97706", "design": "7C3AED", "safety": "DC2626"}
CALLOUT_LABEL_COLOR = {"note": BLUE, "gotcha": AMBER, "design": PURPLE, "safety": RED}


def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_left_border(cell, hex_color, sz=36):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), hex_color)
    borders.append(left)
    tcPr.append(borders)


def set_cell_margins(cell, top=120, bottom=120, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def add_runs_markup(paragraph, text, base_size=10.5, color=None, font="Calibri"):
    tokens = re.split(r"(\*\*.+?\*\*|\*.+?\*)", text)
    for tok in tokens:
        if not tok:
            continue
        bold = tok.startswith("**") and tok.endswith("**")
        italic = (not bold) and tok.startswith("*") and tok.endswith("*")
        clean = tok[2:-2] if bold else (tok[1:-1] if italic else tok)
        run = paragraph.add_run(clean)
        run.font.name = font
        run.font.size = Pt(base_size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
    return paragraph


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(10 if level == 1 else 6)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(19 if level == 1 else 13.5)
    run.font.bold = True
    run.font.color.rgb = BLUE if level == 1 else EMERALD
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        bottom = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "2563EB")
        bottom.append(b)
        pPr.append(bottom)
    return p


def add_para(doc, text, size=10.6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_runs_markup(p, text, base_size=size, color=INK)
    return p


def add_callout(doc, style, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_background(cell, CALLOUT_FILL[style])
    set_cell_left_border(cell, CALLOUT_BAR[style], sz=48)
    set_cell_margins(cell)
    cell.paragraphs[0].text = ""
    lp = cell.paragraphs[0]
    lrun = lp.add_run(label.upper())
    lrun.font.name = "Calibri"
    lrun.font.size = Pt(9.5)
    lrun.font.bold = True
    lrun.font.color.rgb = CALLOUT_LABEL_COLOR[style]
    lp.paragraph_format.space_after = Pt(4)
    tp = cell.add_paragraph()
    add_runs_markup(tp, text, base_size=10.3, color=INK)
    tp.paragraph_format.space_after = Pt(0)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return table


def add_image(doc, path, caption, width_in):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(os.path.join(ASSETS_DIR, path), width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        crun = cp.add_run(caption)
        crun.font.name = "Calibri"
        crun.font.size = Pt(9)
        crun.italic = True
        crun.font.color.rgb = MUTED


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_background(cell, "1A1F2B")
    set_cell_margins(cell)
    cell.paragraphs[0].text = ""
    for i, line in enumerate(text.split("\n")):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = para.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9.3)
        run.font.color.rgb = RGBColor(0xA5, 0xD8, 0xFF)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_bullets(doc, items, numbered=False):
    for item in items:
        p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        p.paragraph_format.space_after = Pt(4)
        add_runs_markup(p, item, base_size=10.5)


def add_glossary(doc, items):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for term, definition in items:
        row = table.add_row()
        c0, c1 = row.cells
        c0.width = Inches(1.8)
        c1.width = Inches(4.8)
        set_cell_margins(c0, left=100, right=100, top=100, bottom=100)
        set_cell_margins(c1, left=100, right=100, top=100, bottom=100)
        set_cell_background(c0, "EFF6FF")
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(term)
        r0.bold = True
        r0.font.name = "Calibri"
        r0.font.size = Pt(10)
        r0.font.color.rgb = BLUE
        p1 = c1.paragraphs[0]
        add_runs_markup(p1, definition, base_size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_table(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        set_cell_background(hdr_cells[i], "1A1F2B")
        set_cell_margins(hdr_cells[i], left=90, right=90, top=70, bottom=70)
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            set_cell_margins(cell, left=90, right=90, top=60, bottom=60)
            p = cell.paragraphs[0]
            add_runs_markup(p, val, base_size=9.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_quiz(doc, items):
    for i, item in enumerate(items, 1):
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(8)
        qrun = qp.add_run(f"{i}. {item['q']}")
        qrun.bold = True
        qrun.font.name = "Calibri"
        qrun.font.size = Pt(10.8)
        qrun.font.color.rgb = INK
        if item["options"]:
            for opt in item["options"]:
                op = doc.add_paragraph(style="List Bullet 2")
                orun = op.add_run(opt)
                orun.font.name = "Calibri"
                orun.font.size = Pt(10)
        ap = doc.add_paragraph()
        ap.paragraph_format.left_indent = Inches(0.25)
        ap.paragraph_format.space_after = Pt(10)
        alabel = ap.add_run("Answer: ")
        alabel.bold = True
        alabel.italic = True
        alabel.font.size = Pt(9.5)
        alabel.font.color.rgb = EMERALD
        atext = ap.add_run(item["answer"])
        atext.italic = True
        atext.font.size = Pt(9.5)
        atext.font.color.rgb = MUTED


def build_cover(doc, block):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(block["title"])
    run.font.name = "Calibri"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(8)
    p2.paragraph_format.space_after = Pt(30)
    run2 = p2.add_run(block["subtitle"])
    run2.font.name = "Calibri"
    run2.font.size = Pt(14)
    run2.italic = True
    run2.font.color.rgb = EMERALD

    for line in block["meta"]:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pm.paragraph_format.space_after = Pt(6)
        rm = pm.add_run(line)
        rm.font.name = "Calibri"
        rm.font.size = Pt(10.5)
        rm.font.color.rgb = MUTED
    doc.add_page_break()


def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    for block in DOCUMENT:
        t = block["type"]
        if t == "cover":
            build_cover(doc, block)
        elif t == "heading":
            add_heading(doc, block["text"], block["level"])
        elif t == "para":
            add_para(doc, block["text"])
        elif t == "callout":
            add_callout(doc, block["style"], block["label"], block["text"])
        elif t == "image":
            add_image(doc, block["path"], block.get("caption", ""), block.get("width_in", 6.0))
        elif t == "code":
            add_code(doc, block["text"])
        elif t == "bullets":
            add_bullets(doc, block["items"])
        elif t == "numbered":
            add_bullets(doc, block["items"], numbered=True)
        elif t == "glossary":
            add_glossary(doc, block["items"])
        elif t == "table":
            add_table(doc, block["header"], block["rows"])
        elif t == "quiz":
            add_quiz(doc, block["items"])
        elif t == "divider":
            doc.add_paragraph("—" * 20).alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif t == "pagebreak":
            doc.add_page_break()
        else:
            raise ValueError(f"Unknown block type: {t}")

    doc.save(OUT_PATH)
    print("Saved:", OUT_PATH)


if __name__ == "__main__":
    main()
