# -*- coding: utf-8 -*-
"""Renders content.DOCUMENT into a .docx family teaching guide."""
import os
import re
import sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(__file__))
from content import DOCUMENT, ASSETS

HERE = os.path.dirname(__file__)
ASSETS_DIR = os.path.join(os.path.dirname(HERE), ASSETS)
OUT_PATH = os.path.join(os.path.dirname(HERE), "Family_Guide_Retail_SQL_Agent.docx")

INK = RGBColor(0x23, 0x2B, 0x3A)
INDIGO = RGBColor(0x2D, 0x3F, 0x6B)
AMBER = RGBColor(0xB5, 0x6E, 0x12)
GREEN = RGBColor(0x2E, 0x5C, 0x44)
MUTED = RGBColor(0x5B, 0x57, 0x48)

CALLOUT_FILL = {"kid": "FCEFD8", "business": "E7F2EC", "note": "E7EBF5"}
CALLOUT_BAR = {"kid": "E0952E", "business": "3F7D5C", "note": "2D3F6B"}
CALLOUT_LABEL_COLOR = {"kid": AMBER, "business": GREEN, "note": INDIGO}


def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_left_border(cell, hex_color, sz=36):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), hex_color)
    borders.append(left)
    tcPr.append(borders)


def set_cell_margins(cell, top=120, bottom=120, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def add_runs_markup(paragraph, text, base_size=11, color=None, font="Calibri"):
    """Parse **bold** and *italic* markup into runs."""
    tokens = re.split(r"(\*\*.+?\*\*|\*.+?\*)", text)
    for tok in tokens:
        if not tok:
            continue
        bold = tok.startswith("**") and tok.endswith("**")
        italic = (not bold) and tok.startswith("*") and tok.endswith("*")
        clean = tok[2:-2] if bold else (tok[1:-1] if italic else tok)
        run = paragraph.add_run(clean)
        run.font.name = font
        run.font.size = Pt(base_size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
    return paragraph


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(10 if level == 1 else 6)
    run = p.add_run(text)
    run.font.name = "Georgia"
    run.font.size = Pt(22 if level == 1 else 15)
    run.font.bold = True
    run.font.color.rgb = INDIGO if level == 1 else GREEN
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        bottom = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "E0952E")
        bottom.append(b)
        pPr.append(bottom)
    return p


def add_para(doc, text, size=11.5, color=None, align=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    if align:
        p.alignment = align
    add_runs_markup(p, text, base_size=size, color=color)
    if italic:
        for r in p.runs:
            r.italic = True
    return p


def add_callout(doc, style, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_background(cell, CALLOUT_FILL[style])
    set_cell_left_border(cell, CALLOUT_BAR[style], sz=48)
    set_cell_margins(cell)
    cell.paragraphs[0].text = ""
    lp = cell.paragraphs[0]
    lrun = lp.add_run(label.upper())
    lrun.font.name = "Calibri"
    lrun.font.size = Pt(10)
    lrun.font.bold = True
    lrun.font.color.rgb = CALLOUT_LABEL_COLOR[style]
    lp.paragraph_format.space_after = Pt(4)
    tp = cell.add_paragraph()
    add_runs_markup(tp, text, base_size=11, color=INK)
    tp.paragraph_format.space_after = Pt(0)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return table


def add_image(doc, path, caption, width_in):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(os.path.join(ASSETS_DIR, path), width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        crun = cp.add_run(caption)
        crun.font.name = "Calibri"
        crun.font.size = Pt(9.5)
        crun.italic = True
        crun.font.color.rgb = MUTED


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_background(cell, "232B3A")
    set_cell_margins(cell)
    cell.paragraphs[0].text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x9B, 0xE8, 0x9B)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_bullets(doc, items, numbered=False):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        p.paragraph_format.space_after = Pt(4)
        add_runs_markup(p, item, base_size=11)


def add_glossary(doc, items):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for term, definition in items:
        row = table.add_row()
        c0, c1 = row.cells
        c0.width = Inches(1.9)
        c1.width = Inches(4.7)
        set_cell_margins(c0, left=100, right=100, top=100, bottom=100)
        set_cell_margins(c1, left=100, right=100, top=100, bottom=100)
        set_cell_background(c0, "F3EFE3")
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(term)
        r0.bold = True
        r0.font.name = "Calibri"
        r0.font.size = Pt(10.5)
        r0.font.color.rgb = INDIGO
        p1 = c1.paragraphs[0]
        add_runs_markup(p1, definition, base_size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_quiz(doc, items):
    for i, item in enumerate(items, 1):
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(8)
        qrun = qp.add_run(f"{i}. {item['q']}")
        qrun.bold = True
        qrun.font.name = "Calibri"
        qrun.font.size = Pt(11.5)
        qrun.font.color.rgb = INK
        if item["options"]:
            for opt in item["options"]:
                op = doc.add_paragraph(style="List Bullet 2")
                orun = op.add_run(opt)
                orun.font.name = "Calibri"
                orun.font.size = Pt(10.5)
        ap = doc.add_paragraph()
        ap.paragraph_format.left_indent = Inches(0.25)
        ap.paragraph_format.space_after = Pt(10)
        alabel = ap.add_run("Answer: ")
        alabel.bold = True
        alabel.italic = True
        alabel.font.size = Pt(10)
        alabel.font.color.rgb = GREEN
        atext = ap.add_run(item["answer"])
        atext.italic = True
        atext.font.size = Pt(10)
        atext.font.color.rgb = MUTED


def build_cover(doc, block):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(block["title"])
    run.font.name = "Georgia"
    run.font.size = Pt(34)
    run.font.bold = True
    run.font.color.rgb = INDIGO

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(8)
    p2.paragraph_format.space_after = Pt(30)
    run2 = p2.add_run(block["subtitle"])
    run2.font.name = "Calibri"
    run2.font.size = Pt(16)
    run2.italic = True
    run2.font.color.rgb = AMBER

    for line in block["meta"]:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pm.paragraph_format.space_after = Pt(6)
        rm = pm.add_run(line)
        rm.font.name = "Calibri"
        rm.font.size = Pt(11.5)
        rm.font.color.rgb = MUTED
    doc.add_page_break()


def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    for block in DOCUMENT:
        t = block["type"]
        if t == "cover":
            build_cover(doc, block)
        elif t == "heading":
            add_heading(doc, block["text"], block["level"])
        elif t == "para":
            add_para(doc, block["text"])
        elif t == "callout":
            add_callout(doc, block["style"], block["label"], block["text"])
        elif t == "image":
            add_image(doc, block["path"], block.get("caption", ""), block.get("width_in", 6.0))
        elif t == "code":
            add_code(doc, block["text"])
        elif t == "bullets":
            add_bullets(doc, block["items"])
        elif t == "numbered":
            add_bullets(doc, block["items"], numbered=True)
        elif t == "glossary":
            add_glossary(doc, block["items"])
        elif t == "quiz":
            add_quiz(doc, block["items"])
        elif t == "divider":
            doc.add_paragraph("—" * 20).alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif t == "pagebreak":
            doc.add_page_break()
        else:
            raise ValueError(f"Unknown block type: {t}")

    doc.save(OUT_PATH)
    print("Saved:", OUT_PATH)


if __name__ == "__main__":
    main()
